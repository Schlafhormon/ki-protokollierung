from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

import main
from export_protocol import parse_summary_sections


EXPORT_PAYLOAD = {
    "format": "txt",
    "metadata": {
        "committee": "Hauptausschuss",
        "date": "2026-06-30",
        "location": "Rathaus",
        "title": "Sitzung Hauptausschuss",
        "participants": ["Alice Beispiel", "Bob Beispiel"],
    },
    "appendix": {
        "include_speaker_list": True,
        "include_transcript_excerpt": True,
        "include_generation_note": True,
    },
    "tops": ["Begrüßung", "Haushalt"],
    "transcript": [
        {"speaker": "SPEAKER_00", "text": "Ich eröffne die Sitzung.", "start": 0, "end": 2},
        {"speaker": "SPEAKER_01", "text": "Der Haushalt wird beraten.", "start": 3, "end": 6},
    ],
    "assignments": [0, 1],
    "speaker_names": {"SPEAKER_00": "Alice", "SPEAKER_01": "Bob"},
    "summaries": {
        "0": "Diskussion:\nDie Sitzung wurde eröffnet.\nBeschluss:\nKeine Beschlüsse.",
        "1": (
            "Diskussion:\nDer Haushalt wurde beraten.\n"
            "Beschluss:\nDer Haushalt wurde empfohlen.\n"
            "Abstimmung:\nEinstimmig.\n"
            "Maßnahmen:\nDie Verwaltung legt Zahlen nach."
        ),
    },
    "summary_reviews": {},
}


def test_parse_summary_sections_extracts_professional_minutes_fields():
    sections = parse_summary_sections(EXPORT_PAYLOAD["summaries"]["1"])

    assert sections["discussion"] == ["Der Haushalt wurde beraten."]
    assert sections["decisions"] == ["Der Haushalt wurde empfohlen."]
    assert sections["votes"] == ["Einstimmig."]
    assert sections["action_items"] == ["Die Verwaltung legt Zahlen nach."]


def test_txt_export_contains_metadata_agenda_sections_and_appendix():
    with TestClient(main.app) as client:
        response = client.post("/api/export", json=EXPORT_PAYLOAD)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/plain")
    content = response.text
    assert "Gremium: Hauptausschuss" in content
    assert "Teilnehmer: Alice Beispiel, Bob Beispiel" in content
    assert "1. Begrüßung" in content
    assert "TOP 2: Haushalt" in content
    assert "Beschluss:" in content
    assert "- Der Haushalt wurde empfohlen." in content
    assert "Abstimmung:" in content
    assert "Maßnahmen/offene Punkte:" in content
    assert "Sprecherliste:" in content
    assert "Transkript-Auszug:" in content
    assert "Bearbeitungs-/Generierungshinweis:" in content


def test_docx_export_has_expected_document_structure():
    payload = {**EXPORT_PAYLOAD, "format": "docx"}

    with TestClient(main.app) as client:
        response = client.post("/api/export", json=payload)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    doc = Document(BytesIO(response.content))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]

    assert "Sitzung Hauptausschuss" in paragraphs
    assert "Tagesordnung" in paragraphs
    assert "TOP 2: Haushalt" in paragraphs
    assert "Beschluss" in paragraphs
    assert "Der Haushalt wurde empfohlen." in paragraphs
    assert "Anhang" in paragraphs

    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "Gremium" in table_text
    assert "Hauptausschuss" in table_text
    assert "Alice Beispiel, Bob Beispiel" in table_text


def test_pdf_export_returns_pdf_bytes_with_attachment_header():
    payload = {**EXPORT_PAYLOAD, "format": "pdf"}

    with TestClient(main.app) as client:
        response = client.post("/api/export", json=payload)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert response.headers["content-disposition"].endswith('.pdf"')
    assert response.content.startswith(b"%PDF")


def test_export_rejects_unknown_format():
    payload = {**EXPORT_PAYLOAD, "format": "xlsx"}

    with TestClient(main.app) as client:
        response = client.post("/api/export", json=payload)

    assert response.status_code == 400
