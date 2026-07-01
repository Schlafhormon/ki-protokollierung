"""Professional protocol export rendering for TXT, DOCX and PDF."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Iterable, Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    ListFlowable,
    ListItem,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from xml.sax.saxutils import escape


ExportFormat = Literal["txt", "docx", "pdf"]

SECTION_LABELS = {
    "discussion": "Diskussion",
    "decisions": "Beschluss",
    "votes": "Abstimmung",
    "action_items": "Maßnahmen/offene Punkte",
    "open_points": "Offene Punkte",
    "uncertainties": "Unsicherheiten",
}

STRUCTURED_KEYS = tuple(SECTION_LABELS.keys())


@dataclass
class ProtocolMetadata:
    committee: str = ""
    date: str = ""
    location: str = ""
    title: str = ""
    participants: list[str] = field(default_factory=list)


@dataclass
class ProtocolTop:
    index: int
    title: str
    discussion: list[str] = field(default_factory=list)
    decisions: list[str] = field(default_factory=list)
    votes: list[str] = field(default_factory=list)
    action_items: list[str] = field(default_factory=list)
    open_points: list[str] = field(default_factory=list)


@dataclass
class ProtocolAppendix:
    include_speaker_list: bool = True
    include_transcript_excerpt: bool = False
    include_generation_note: bool = True
    transcript_excerpt_limit: int = 20


@dataclass
class TranscriptLine:
    speaker: str
    text: str
    start: float = 0.0
    end: float = 0.0


@dataclass
class ProtocolDocument:
    metadata: ProtocolMetadata
    tops: list[ProtocolTop]
    agenda: list[str]
    speakers: list[str]
    transcript: list[TranscriptLine]
    appendix: ProtocolAppendix


def normalize_participants(value: Iterable[str] | str | None) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        raw_items = re.split(r"[\n;,]+", value)
    else:
        raw_items = [str(item) for item in value]
    return [item.strip() for item in raw_items if item and item.strip()]


def format_timestamp(seconds: float | int | None) -> str:
    seconds = max(0, int(seconds or 0))
    mins, secs = divmod(seconds, 60)
    hours, mins = divmod(mins, 60)
    if hours:
        return f"{hours}:{mins:02d}:{secs:02d}"
    return f"{mins}:{secs:02d}"


def parse_summary_sections(summary: str | None) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {key: [] for key in STRUCTURED_KEYS}
    if not summary or not summary.strip():
        return sections

    label_to_key = {
        "diskussion": "discussion",
        "beschluss": "decisions",
        "beschluesse": "decisions",
        "beschlusse": "decisions",
        "abstimmung": "votes",
        "massnahmen": "action_items",
        "maßnahmen": "action_items",
        "massnahmen/offene punkte": "action_items",
        "maßnahmen/offene punkte": "action_items",
        "offene punkte": "open_points",
        "unsicherheiten": "uncertainties",
    }
    current_key = "discussion"
    found_heading = False

    for raw_line in summary.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        heading_match = re.match(r"^([A-Za-zÄÖÜäöüß /-]+):\s*(.*)$", line)
        if heading_match:
            heading = heading_match.group(1).strip().lower()
            if heading in label_to_key:
                current_key = label_to_key[heading]
                found_heading = True
                remainder = heading_match.group(2).strip()
                if remainder:
                    sections[current_key].append(_strip_bullet(remainder))
                continue

        sections[current_key].append(_strip_bullet(line))

    if not found_heading and not sections["discussion"]:
        sections["discussion"] = [_strip_bullet(summary.strip())]
    return sections


def build_protocol_document(
    *,
    metadata: ProtocolMetadata,
    tops: list[str],
    summaries: dict[int, str],
    summary_reviews: dict[int, dict] | None = None,
    transcript: list[TranscriptLine] | None = None,
    speaker_names: dict[str, str] | None = None,
    appendix: ProtocolAppendix | None = None,
) -> ProtocolDocument:
    summary_reviews = summary_reviews or {}
    transcript = transcript or []
    speaker_names = speaker_names or {}
    appendix = appendix or ProtocolAppendix()

    agenda = [top.strip() or f"TOP {index + 1}" for index, top in enumerate(tops)]
    protocol_tops: list[ProtocolTop] = []

    for index, title in enumerate(agenda):
        editable_summary = summaries.get(index)
        structured = (summary_reviews.get(index) or {}).get("structured")
        if editable_summary and editable_summary.strip():
            sections = parse_summary_sections(editable_summary)
        elif structured:
            sections = {
                key: [str(item).strip() for item in structured.get(key, []) if str(item).strip()]
                for key in STRUCTURED_KEYS
            }
        else:
            sections = parse_summary_sections(editable_summary)

        protocol_tops.append(
            ProtocolTop(
                index=index + 1,
                title=title,
                discussion=sections["discussion"],
                decisions=sections["decisions"],
                votes=sections["votes"],
                action_items=sections["action_items"],
                open_points=sections["open_points"],
            )
        )

    speaker_set = []
    for line in transcript:
        display = speaker_names.get(line.speaker, line.speaker)
        if display not in speaker_set:
            speaker_set.append(display)

    return ProtocolDocument(
        metadata=metadata,
        tops=protocol_tops,
        agenda=agenda,
        speakers=speaker_set,
        transcript=transcript,
        appendix=appendix,
    )


def render_protocol(document: ProtocolDocument, export_format: ExportFormat) -> bytes:
    if export_format == "txt":
        return render_txt(document).encode("utf-8")
    if export_format == "docx":
        return render_docx(document)
    if export_format == "pdf":
        return render_pdf(document)
    raise ValueError(f"Unsupported export format: {export_format}")


def render_txt(document: ProtocolDocument) -> str:
    lines: list[str] = ["SITZUNGSPROTOKOLL", "=" * 60, ""]
    lines.extend(_metadata_text_lines(document.metadata))
    lines.append("")
    lines.append("Tagesordnung")
    lines.append("-" * 60)
    for index, top in enumerate(document.agenda, start=1):
        lines.append(f"{index}. {top}")
    lines.append("")

    for top in document.tops:
        lines.append(f"TOP {top.index}: {top.title}")
        lines.append("-" * 60)
        _append_text_section(lines, "Diskussion", top.discussion)
        _append_text_section(lines, "Beschluss", top.decisions)
        _append_text_section(lines, "Abstimmung", top.votes)
        _append_text_section(lines, "Maßnahmen/offene Punkte", top.action_items + top.open_points)
        lines.append("")

    _append_text_appendix(lines, document)
    return "\n".join(lines).rstrip() + "\n"


def render_docx(document: ProtocolDocument) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    doc.styles["Title"].font.name = "Arial"
    doc.styles["Title"].font.size = Pt(20)
    doc.styles["Heading 1"].font.name = "Arial"
    doc.styles["Heading 2"].font.name = "Arial"

    title = document.metadata.title or "Sitzungsprotokoll"
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in _metadata_rows(document.metadata):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value or "-"
    doc.add_paragraph()

    doc.add_heading("Tagesordnung", level=1)
    for index, agenda_item in enumerate(document.agenda, start=1):
        doc.add_paragraph(f"{index}. {agenda_item}", style="List Number")

    for top in document.tops:
        doc.add_heading(f"TOP {top.index}: {top.title}", level=1)
        _add_docx_section(doc, "Diskussion", top.discussion)
        _add_docx_section(doc, "Beschluss", top.decisions)
        _add_docx_section(doc, "Abstimmung", top.votes)
        _add_docx_section(
            doc,
            "Maßnahmen/offene Punkte",
            top.action_items + top.open_points,
        )

    _add_docx_appendix(doc, document)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def render_pdf(document: ProtocolDocument) -> bytes:
    output = io.BytesIO()
    doc = BaseDocTemplate(
        output,
        pagesize=A4,
        leftMargin=2.0 * cm,
        rightMargin=2.0 * cm,
        topMargin=2.0 * cm,
        bottomMargin=1.8 * cm,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="protocol", frames=[frame], onPage=_pdf_footer)])

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="ProtocolTitle", parent=styles["Title"], alignment=TA_CENTER, fontSize=18, leading=22, spaceAfter=14))
    styles.add(ParagraphStyle(name="SectionTitle", parent=styles["Heading1"], fontSize=13, leading=16, spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name="TopTitle", parent=styles["Heading1"], fontSize=14, leading=17, spaceBefore=16, spaceAfter=8))
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8.5, leading=11))

    story = [Paragraph(_pdf_text(document.metadata.title or "Sitzungsprotokoll"), styles["ProtocolTitle"])]

    rows = [[Paragraph(f"<b>{_pdf_text(label)}</b>", styles["BodyText"]), Paragraph(_pdf_text(value or "-"), styles["BodyText"])] for label, value in _metadata_rows(document.metadata)]
    table = Table(rows, colWidths=[4.2 * cm, 10.8 * cm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F4F6")),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D1D5DB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([table, Spacer(1, 10), Paragraph("Tagesordnung", styles["SectionTitle"])])
    story.append(_pdf_list([f"{index}. {top}" for index, top in enumerate(document.agenda, start=1)], styles["BodyText"]))

    for top in document.tops:
        story.append(Paragraph(_pdf_text(f"TOP {top.index}: {top.title}"), styles["TopTitle"]))
        _append_pdf_section(story, styles, "Diskussion", top.discussion)
        _append_pdf_section(story, styles, "Beschluss", top.decisions)
        _append_pdf_section(story, styles, "Abstimmung", top.votes)
        _append_pdf_section(story, styles, "Maßnahmen/offene Punkte", top.action_items + top.open_points)

    _append_pdf_appendix(story, styles, document)
    doc.build(story)
    return output.getvalue()


def _strip_bullet(text: str) -> str:
    return re.sub(r"^[-*•]\s*", "", text).strip()


def _metadata_rows(metadata: ProtocolMetadata) -> list[tuple[str, str]]:
    return [
        ("Gremium", metadata.committee),
        ("Datum", metadata.date),
        ("Ort", metadata.location),
        ("Sitzungstitel", metadata.title),
        ("Teilnehmer", ", ".join(metadata.participants)),
    ]


def _metadata_text_lines(metadata: ProtocolMetadata) -> list[str]:
    return [f"{label}: {value or '-'}" for label, value in _metadata_rows(metadata)]


def _append_text_section(lines: list[str], label: str, items: list[str]) -> None:
    lines.append(f"{label}:")
    if items:
        lines.extend(f"- {item}" for item in items)
    else:
        lines.append("- Keine Angabe.")
    lines.append("")


def _append_text_appendix(lines: list[str], document: ProtocolDocument) -> None:
    appendix_lines: list[str] = []
    if document.appendix.include_speaker_list and document.speakers:
        appendix_lines.extend(["Sprecherliste:", *[f"- {speaker}" for speaker in document.speakers], ""])
    if document.appendix.include_transcript_excerpt and document.transcript:
        appendix_lines.append("Transkript-Auszug:")
        for line in document.transcript[: document.appendix.transcript_excerpt_limit]:
            appendix_lines.append(f"- [{format_timestamp(line.start)}] {line.speaker}: {line.text}")
        appendix_lines.append("")
    if document.appendix.include_generation_note:
        appendix_lines.extend(["Bearbeitungs-/Generierungshinweis:", "- Automatisch erzeugter Entwurf; fachlich und rechtlich zu prüfen.", ""])
    if appendix_lines:
        lines.extend(["Anhang", "=" * 60, *appendix_lines])


def _add_docx_section(doc: Document, label: str, items: list[str]) -> None:
    doc.add_heading(label, level=2)
    if items:
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("Keine Angabe.")


def _add_docx_appendix(doc: Document, document: ProtocolDocument) -> None:
    if not (
        document.appendix.include_speaker_list
        or document.appendix.include_transcript_excerpt
        or document.appendix.include_generation_note
    ):
        return

    doc.add_page_break()
    doc.add_heading("Anhang", level=1)
    if document.appendix.include_speaker_list and document.speakers:
        doc.add_heading("Sprecherliste", level=2)
        for speaker in document.speakers:
            doc.add_paragraph(speaker, style="List Bullet")
    if document.appendix.include_transcript_excerpt and document.transcript:
        doc.add_heading("Transkript-Auszug", level=2)
        for line in document.transcript[: document.appendix.transcript_excerpt_limit]:
            doc.add_paragraph(f"[{format_timestamp(line.start)}] {line.speaker}: {line.text}")
    if document.appendix.include_generation_note:
        doc.add_heading("Bearbeitungs-/Generierungshinweis", level=2)
        doc.add_paragraph("Automatisch erzeugter Entwurf; fachlich und rechtlich zu prüfen.")


def _append_pdf_section(story: list, styles, label: str, items: list[str]) -> None:
    story.append(Paragraph(_pdf_text(label), styles["SectionTitle"]))
    story.append(_pdf_list(items or ["Keine Angabe."], styles["BodyText"]))


def _append_pdf_appendix(story: list, styles, document: ProtocolDocument) -> None:
    if not (
        document.appendix.include_speaker_list
        or document.appendix.include_transcript_excerpt
        or document.appendix.include_generation_note
    ):
        return
    story.append(Paragraph("Anhang", styles["TopTitle"]))
    if document.appendix.include_speaker_list and document.speakers:
        story.append(Paragraph("Sprecherliste", styles["SectionTitle"]))
        story.append(_pdf_list(document.speakers, styles["BodyText"]))
    if document.appendix.include_transcript_excerpt and document.transcript:
        story.append(Paragraph("Transkript-Auszug", styles["SectionTitle"]))
        excerpt = [
            f"[{format_timestamp(line.start)}] {line.speaker}: {line.text}"
            for line in document.transcript[: document.appendix.transcript_excerpt_limit]
        ]
        story.append(_pdf_list(excerpt, styles["Small"]))
    if document.appendix.include_generation_note:
        story.append(Paragraph("Bearbeitungs-/Generierungshinweis", styles["SectionTitle"]))
        story.append(Paragraph("Automatisch erzeugter Entwurf; fachlich und rechtlich zu prüfen.", styles["BodyText"]))


def _pdf_list(items: list[str], style) -> ListFlowable:
    return ListFlowable(
        [ListItem(Paragraph(_pdf_text(item), style), leftIndent=10) for item in items],
        bulletType="bullet",
        leftIndent=16,
        bulletFontSize=6,
    )


def _pdf_text(value: str) -> str:
    return escape(str(value)).replace("\n", "<br/>")


def _pdf_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    canvas.drawRightString(19 * cm, 1.1 * cm, f"Seite {doc.page}")
    canvas.drawString(2 * cm, 1.1 * cm, datetime.now().strftime("Export: %d.%m.%Y"))
    canvas.restoreState()
